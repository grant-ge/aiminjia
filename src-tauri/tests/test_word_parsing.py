#!/usr/bin/env python3
"""
Test Word document parsing with various edge cases.
Creates test .docx files and validates the parsing logic.
"""

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# Test output directory
TEST_DIR = Path(__file__).parent / "fixtures" / "word"
TEST_DIR.mkdir(parents=True, exist_ok=True)


def create_test_doc_merged_cells():
    """Test case 1: Table with merged cells (the bug we just fixed)"""
    doc = Document()
    doc.add_heading('Test: Merged Cells', level=1)

    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'

    # Header row with merged cells
    header_cells = table.rows[0].cells
    header_cells[0].merge(header_cells[4])  # Merge all 5 columns into 1
    header_cells[0].text = 'Merged Header'

    # Data rows with normal 5 columns
    for i, row in enumerate(table.rows[1:], 1):
        for j, cell in enumerate(row.cells):
            cell.text = f'R{i}C{j}'

    path = TEST_DIR / "merged_cells.docx"
    doc.save(path)
    print(f"✓ Created: {path}")
    return path


def create_test_doc_multiple_tables():
    """Test case 2: Multiple tables with different column counts"""
    doc = Document()
    doc.add_heading('Test: Multiple Tables', level=1)

    # Table 1: 3 columns
    doc.add_paragraph('Table 1: Employee Data')
    table1 = doc.add_table(rows=3, cols=3)
    table1.style = 'Light Grid Accent 1'
    headers1 = ['Name', 'Age', 'Department']
    for i, header in enumerate(headers1):
        table1.rows[0].cells[i].text = header
    table1.rows[1].cells[0].text = 'Alice'
    table1.rows[1].cells[1].text = '30'
    table1.rows[1].cells[2].text = 'Engineering'

    doc.add_paragraph()  # Spacer

    # Table 2: 5 columns (different structure)
    doc.add_paragraph('Table 2: Sales Data')
    table2 = doc.add_table(rows=3, cols=5)
    table2.style = 'Light Grid Accent 1'
    headers2 = ['Product', 'Q1', 'Q2', 'Q3', 'Q4']
    for i, header in enumerate(headers2):
        table2.rows[0].cells[i].text = header
    table2.rows[1].cells[0].text = 'Widget'
    for i in range(1, 5):
        table2.rows[1].cells[i].text = str(100 * i)

    path = TEST_DIR / "multiple_tables.docx"
    doc.save(path)
    print(f"✓ Created: {path}")
    return path


def create_test_doc_text_only():
    """Test case 3: Pure text document (no tables)"""
    doc = Document()
    doc.add_heading('Test: Text Only', level=1)
    doc.add_paragraph('This is a plain text document.')
    doc.add_paragraph('It contains multiple paragraphs.')
    doc.add_paragraph('But no tables at all.')
    doc.add_heading('Section 2', level=2)
    doc.add_paragraph('More text content here.')

    path = TEST_DIR / "text_only.docx"
    doc.save(path)
    print(f"✓ Created: {path}")
    return path


def create_test_doc_empty_cells():
    """Test case 4: Table with empty cells"""
    doc = Document()
    doc.add_heading('Test: Empty Cells', level=1)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header
    headers = ['Name', 'Value', 'Notes']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Data with empty cells
    table.rows[1].cells[0].text = 'Item 1'
    table.rows[1].cells[1].text = '100'
    # cells[2] is empty

    table.rows[2].cells[0].text = 'Item 2'
    # cells[1] is empty
    table.rows[2].cells[2].text = 'Some note'

    # Row 3: all empty

    path = TEST_DIR / "empty_cells.docx"
    doc.save(path)
    print(f"✓ Created: {path}")
    return path


def create_test_doc_complex_merge():
    """Test case 5: Complex merged cell patterns"""
    doc = Document()
    doc.add_heading('Test: Complex Merge', level=1)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'

    # Row 0: Normal header
    headers = ['Col1', 'Col2', 'Col3', 'Col4']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Row 1: Merge first two cells
    table.rows[1].cells[0].merge(table.rows[1].cells[1])
    table.rows[1].cells[0].text = 'Merged 1-2'
    table.rows[1].cells[2].text = 'C'
    table.rows[1].cells[3].text = 'D'

    # Row 2: Merge last two cells
    table.rows[2].cells[0].text = 'A'
    table.rows[2].cells[1].text = 'B'
    table.rows[2].cells[2].merge(table.rows[2].cells[3])
    table.rows[2].cells[2].text = 'Merged 3-4'

    # Row 3: Normal
    for i in range(4):
        table.rows[3].cells[i].text = f'R3C{i}'

    # Row 4: Merge all
    table.rows[4].cells[0].merge(table.rows[4].cells[3])
    table.rows[4].cells[0].text = 'Full row merge'

    path = TEST_DIR / "complex_merge.docx"
    doc.save(path)
    print(f"✓ Created: {path}")
    return path


def test_parse_logic(docx_path):
    """
    Test the actual parsing logic (mimics parser.rs parse_word function).
    This is the Python code that runs inside Rust's parser.
    """
    print(f"\n--- Testing: {docx_path.name} ---")

    try:
        doc = Document(docx_path)

        # Extract tables (same logic as parser.rs)
        # Strategy: extract the largest table (by row count)
        best_table = None
        best_row_count = 0

        for table in doc.tables:
            if len(table.rows) > best_row_count:
                best_row_count = len(table.rows)
                best_table = table

        all_rows = []
        header = None

        if best_table is not None:
            for i, row in enumerate(best_table.rows):
                # Deduplicate merged cells
                seen = set()
                unique_cells = []
                for cell in row.cells:
                    cell_id = id(cell)
                    if cell_id not in seen:
                        seen.add(cell_id)
                        unique_cells.append(cell.text.strip())

                # Skip rows with only 1 unique cell (fully merged rows)
                if len(unique_cells) == 1:
                    continue

                if header is None and any(unique_cells):
                    header = unique_cells
                    print(f"  Header: {header} (len={len(header)})")
                elif header is not None:
                    # Align column count
                    n = len(header)
                    if len(unique_cells) < n:
                        unique_cells.extend([''] * (n - len(unique_cells)))
                    elif len(unique_cells) > n:
                        unique_cells = unique_cells[:n]
                    all_rows.append(unique_cells)

        if header and all_rows:
            print(f"  ✓ Extracted table: {len(all_rows)} rows, {len(header)} columns")
            print(f"  Sample rows:")
            for i, row in enumerate(all_rows[:3], 1):
                print(f"    Row {i}: {row}")
            return {"type": "table", "rows": len(all_rows), "cols": len(header)}
        else:
            # Extract text
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            print(f"  ✓ Extracted text: {len(paragraphs)} paragraphs")
            print(f"  Preview: {paragraphs[0][:50]}..." if paragraphs else "  (empty)")
            return {"type": "text", "paragraphs": len(paragraphs)}

    except Exception as e:
        print(f"  ✗ FAILED: {e}")
        import traceback
        traceback.print_exc()
        return {"type": "error", "error": str(e)}


def main():
    print("Creating test Word documents...\n")

    test_files = [
        create_test_doc_merged_cells(),
        create_test_doc_multiple_tables(),
        create_test_doc_text_only(),
        create_test_doc_empty_cells(),
        create_test_doc_complex_merge(),
    ]

    print("\n" + "="*60)
    print("Testing parsing logic...")
    print("="*60)

    results = []
    for path in test_files:
        result = test_parse_logic(path)
        results.append({"file": path.name, "result": result})

    print("\n" + "="*60)
    print("Summary:")
    print("="*60)

    failed = [r for r in results if r["result"]["type"] == "error"]
    if failed:
        print(f"\n✗ {len(failed)} test(s) FAILED:")
        for r in failed:
            print(f"  - {r['file']}: {r['result']['error']}")
        sys.exit(1)
    else:
        print(f"\n✓ All {len(results)} tests PASSED")
        for r in results:
            print(f"  - {r['file']}: {r['result']['type']}")


if __name__ == "__main__":
    main()
