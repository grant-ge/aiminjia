"""Parse a DOCX template and emit chapter outline + style hints as JSON.

Usage (from execute_python tool):
    import os
    os.environ["TEMPLATE_PATH"] = "/abs/path/to/template.docx"
    exec(open("references/parse_template.py").read())

For PDF templates, fall back to load_file's text extraction in the skill flow;
this script handles DOCX only.
"""

import json
import os
import sys

from docx import Document

template_path = os.environ.get("TEMPLATE_PATH")
if not template_path:
    print(json.dumps({"error": "TEMPLATE_PATH env var not set"}))
    sys.exit(1)

if not os.path.exists(template_path):
    print(json.dumps({"error": f"file not found: {template_path}"}))
    sys.exit(1)

doc = Document(template_path)

chapters = []
for para in doc.paragraphs:
    style_name = para.style.name if para.style else ""
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.replace("Heading ", "").strip())
        except (ValueError, AttributeError):
            level = 1
        text = para.text.strip()
        if text:
            chapters.append({"level": level, "title": text})

style_hint = {
    "title_font": None,
    "body_font": None,
    "title_color": None,
}

for para in doc.paragraphs:
    if para.style and para.style.name == "Heading 1" and para.runs:
        run = para.runs[0]
        if run.font.name:
            style_hint["title_font"] = run.font.name
        if run.font.color and run.font.color.rgb:
            style_hint["title_color"] = "#" + str(run.font.color.rgb)
        break

for para in doc.paragraphs:
    if para.style and para.style.name == "Normal" and para.runs:
        run = para.runs[0]
        if run.font.name:
            style_hint["body_font"] = run.font.name
            break

print(json.dumps(
    {"chapters": chapters, "style_hint": style_hint},
    ensure_ascii=False,
))
